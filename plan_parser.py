#!/usr/bin/env python3
"""
Plan Parser - Extract and assign roles from the dismissal plan
"""

from docx import Document
from pathlib import Path
import re
from typing import Dict, List, Tuple

class PlanParser:
    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.supervisors = {}
        self.teacher_assignments = []
        self.gate_assignments = {}
        self.daily_assignments = {}  # New: day-specific teacher assignments
        
    def parse_plan(self):
        """Parse the entire plan and extract all assignments"""
        self._extract_supervisors()
        self._extract_teacher_assignments()
        self._extract_gate_assignments()
        self._extract_daily_teacher_assignments()  # New method
        
    def _extract_supervisors(self):
        """Extract daily supervisors from paragraphs"""
        current_day = None
        current_supervisor = None
        current_subjects = []
        
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Look for day patterns
            day_match = re.search(r'Day:\s*(\w+)', text, re.IGNORECASE)
            if day_match:
                current_day = day_match.group(1).capitalize()
                # Also extract subjects from the same line
                current_subjects = self._extract_subjects(text)
                continue
            
            # Look for supervisor patterns
            supervisor_match = re.search(r'Supervis(?:or|ion):\s*([^,\n]+)', text)
            if supervisor_match:
                current_supervisor = supervisor_match.group(1).strip()
                
                # If we have both day and supervisor, store them
                if current_day and current_supervisor:
                    self.supervisors[current_day] = {
                        'name': current_supervisor,
                        'subjects': current_subjects
                    }
                    # Reset for next day
                    current_day = None
                    current_supervisor = None
                    current_subjects = []
            
            # Look for team information
            if 'Team:' in text and current_day:
                team_subjects = self._extract_subjects(text)
                if team_subjects:
                    current_subjects.extend(team_subjects)
    
    def _extract_subjects(self, text: str) -> List[str]:
        """Extract subjects from text"""
        # Common subject patterns
        subjects = []
        subject_patterns = [
            r'ELA', r'German', r'SSA', r'Music', r'French',
            r'Arabic', r'SSE', r'Computer', r'Math', r'Science'
        ]
        
        for pattern in subject_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                subjects.append(pattern)
        
        return subjects
    
    def _extract_teacher_assignments(self):
        """Extract teacher assignments from tables"""
        for table in self.doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                
                if len(cells) >= 3:  # Number, Name, Role
                    try:
                        # Try to parse as: Number | Name | Role
                        if cells[0].isdigit():
                            teacher_id = int(cells[0])
                            teacher_name = cells[1]
                            role = cells[2]
                            
                            self.teacher_assignments.append({
                                'id': teacher_id,
                                'name': teacher_name,
                                'role': role
                            })
                    except (ValueError, IndexError):
                        continue
    
    def _extract_gate_assignments(self):
        """Extract gate assignments from tables"""
        for table in self.doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                
                # Look for gate-related rows
                if any('gate' in cell.lower() for cell in cells):
                    if len(cells) >= 2:
                        # First row might be headers, subsequent rows are assignments
                        gates = [cell for cell in cells if 'gate' in cell.lower()]
                        staff = [cell for cell in cells if 'gate' not in cell.lower() and cell]
                        
                        for i, gate in enumerate(gates):
                            if i < len(staff):
                                self.gate_assignments[gate] = staff[i]
    
    def _extract_daily_teacher_assignments(self):
        """Extract day-specific teacher assignments by matching tables to day declarations"""
        # Get all elements in document order (paragraphs and tables)
        elements = []
        for element in self.doc.element.body:
            if element.tag.endswith('p'):  # paragraph
                para = None
                for p in self.doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para and para.text.strip():
                    elements.append(('paragraph', para.text.strip()))
            elif element.tag.endswith('tbl'):  # table
                table = None
                for t in self.doc.tables:
                    if t._element == element:
                        table = t
                        break
                if table:
                    elements.append(('table', table))
        
        # Map days to their information
        day_info = {}
        current_day = None
        current_supervisor = None
        
        # First pass: collect day and supervisor information
        for element_type, content in elements:
            if element_type == 'paragraph':
                # Check for day information
                if 'Day:' in content:
                    day_match = re.search(r'Day:\s*(\w+)', content)
                    if day_match:
                        current_day = day_match.group(1)
                        if current_day not in day_info:
                            day_info[current_day] = {'supervisor': None, 'team': None}
                
                # Check for supervisor information
                elif 'Supervis' in content:
                    supervisor_match = re.search(r'Supervis(?:or|ion):\s*([^,\n]+)', content)
                    if supervisor_match:
                        current_supervisor = supervisor_match.group(1).strip()
                        # If we have a current day, assign supervisor to it
                        if current_day and current_day in day_info:
                            day_info[current_day]['supervisor'] = current_supervisor
                        # Otherwise, this might be for the next day we encounter
                        else:
                            # Look ahead for the next day
                            temp_supervisor = current_supervisor
                
                # Check for team information
                elif 'Team:' in content and current_day:
                    team_match = re.search(r'Team:\s*(.+)', content)
                    if team_match:
                        current_team = team_match.group(1).strip()
                        day_info[current_day]['team'] = current_team
        
        # Handle the special case where Sunday supervisor comes before Sunday day declaration
        # Based on document structure: Supervisor: Hany Amin, then ELA + German + SSA + Music Day: Sunday
        if 'Sunday' in day_info and not day_info['Sunday']['supervisor']:
            # Look for Hany Amin supervisor before Sunday
            for element_type, content in elements:
                if element_type == 'paragraph' and 'Supervisor: Hany Amin' in content:
                    day_info['Sunday']['supervisor'] = 'Hany Amin'
                    break
        
        # Second pass: assign tables to days
        table_assignments = []
        current_day = None
        
        for element_type, content in elements:
            if element_type == 'paragraph':
                if 'Day:' in content:
                    day_match = re.search(r'Day:\s*(\w+)', content)
                    if day_match:
                        current_day = day_match.group(1)
            elif element_type == 'table':
                if current_day:
                    table_assignments.append((current_day, content))
                    current_day = None  # Reset to avoid assigning next table to same day
                else:
                    # This might be the Sunday table (first table without preceding day)
                    if len(table_assignments) == 0:
                        table_assignments.append(('Sunday', content))
        
        # Extract teachers from assigned tables
        for day, table in table_assignments:
            if day not in self.daily_assignments:
                self.daily_assignments[day] = {
                    'supervisor': day_info.get(day, {}).get('supervisor'),
                    'team': day_info.get(day, {}).get('team'),
                    'teachers': []
                }
            
            teachers = []
            for row in table.rows[1:]:  # Skip header row
                row_cells = [cell.text.strip() for cell in row.cells]
                
                # Skip supervision rows and empty rows
                if (len(row_cells) >= 2 and 
                    'Supervision' not in row_cells[1] and
                    row_cells[0].isdigit() and
                    row_cells[1]):  # Has teacher name
                    
                    teacher_name = row_cells[1]
                    teacher_role = row_cells[2] if len(row_cells) > 2 else ""
                    
                    teachers.append({
                        'name': teacher_name,
                        'role': teacher_role
                    })
            
            self.daily_assignments[day]['teachers'] = teachers

    def get_assignments_summary(self) -> Dict:
        """Get a summary of all assignments"""
        return {
            'supervisors': self.supervisors,
            'teachers': self.teacher_assignments,
            'gates': self.gate_assignments,
            'daily_assignments': self.daily_assignments  # Include daily assignments
        }
    
    def print_assignments(self):
        """Print all assignments in a readable format"""
        print("=== DISMISSAL PLAN ASSIGNMENTS ===\n")
        
        print("📅 DAILY SUPERVISORS:")
        for day, info in self.supervisors.items():
            subjects = ', '.join(info['subjects']) if info['subjects'] else 'No subjects specified'
            print(f"  {day}: {info['name']} ({subjects})")
        
        print(f"\n=== DAILY TEACHER ASSIGNMENTS ===")
        for day, info in self.daily_assignments.items():
            print(f"\n📅 {day}")
            if info['supervisor']:
                print(f"  👨‍🏫 Supervisor: {info['supervisor']}")
            if info['team']:
                print(f"  👥 Team: {info['team']}")
            if info['teachers']:
                print(f"  👩‍🏫 Teachers ({len(info['teachers'])}):") 
                for teacher in info['teachers']:
                    print(f"    - {teacher['name']}")
            else:
                print("  No teachers assigned")
        
        print(f"\n👥 GENERAL TEACHER ASSIGNMENTS ({len(self.teacher_assignments)} teachers):")
        for teacher in self.teacher_assignments:
            print(f"  {teacher['id']:2d}. {teacher['name']:<20} → {teacher['role']}")
        
        print(f"\n🚪 GATE ASSIGNMENTS:")
        for gate, staff in self.gate_assignments.items():
            print(f"  {gate}: {staff}")

if __name__ == "__main__":
    # Parse the uploaded plan
    doc_path = "resources/Dismissal_plan_2024-2025_2_20251001_135735.docx"
    
    if Path(doc_path).exists():
        parser = PlanParser(doc_path)
        parser.parse_plan()
        parser.print_assignments()
        
        # Save assignments to a JSON file for later use
        import json
        assignments = parser.get_assignments_summary()
        
        with open("data/parsed_assignments.json", "w", encoding="utf-8") as f:
            json.dump(assignments, f, indent=2, ensure_ascii=False)
        
        print(f"\n✅ Assignments saved to data/parsed_assignments.json")
        print(f"📊 Summary: {len(assignments['supervisors'])} supervisors, {len(assignments['teachers'])} teachers, {len(assignments['gates'])} gates, {len(assignments['daily_assignments'])} daily plans")
    else:
        print(f"❌ Document not found: {doc_path}")