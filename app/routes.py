from __future__ import annotations

import os
from datetime import date, datetime
from pathlib import Path
import pytz

from docx import Document
from flask import (
    Blueprint,
    current_app,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    url_for,
    send_from_directory,
)
from werkzeug.utils import secure_filename
from xlrd import XL_CELL_NUMBER, open_workbook

from . import db
from .models import AttendanceRecord, DutyAssignment, DutyPlan, DutySection, Teacher
from .email_service import (
    send_bulk_attendance_notifications,
    send_attendance_notification,
    send_absence_warning_email,
    send_absence_deduction_email,
    send_late_warning_email,
    send_late_deduction_email,
    send_new_system_announcement,
    ADMIN_CC_EMAILS,
    SMTP_SERVER,
    SMTP_PORT,
    EMAIL_ADDRESS,
    EMAIL_PASSWORD
)

main_bp = Blueprint("main", __name__)


@main_bp.app_errorhandler(404)
def not_found(error):  # pragma: no cover - template driven
    return render_template("404.html"), 404


@main_bp.route("/")
def index():
    return redirect(url_for("main.today_plan"))


@main_bp.route("/offline-confirmation")
def offline_confirmation():
    return render_template("offline_confirmation.html")


@main_bp.route("/send-notifications", methods=["GET", "POST"])
def send_notifications():
    """
    Process today's attendance and send appropriate notifications:
    - Absent teachers: warning 1, 2 or 3 (with 1/5 day deduction on 3rd, then reset)
    - Late teachers: warning 1, 2, 3 or 4 (with half day deduction on 4th, then reset)
    """
    today = date.today()
    today_day_name = today.strftime("%A")
    
    # Get today's plan
    plan = DutyPlan.query.filter_by(day_of_week=today_day_name).first()
    
    if not plan:
        if request.headers.get('Accept') == 'application/json':
            return jsonify({'success': False, 'message': f'No plan found for {today_day_name}'}), 404
        flash(f'No plan found for {today_day_name}', 'warning')
        return redirect(url_for('main.today_plan'))
    
    # Get today's attendance records
    attendance_records = db.session.query(
        AttendanceRecord, DutyAssignment, DutySection, Teacher
    ).join(
        DutyAssignment, AttendanceRecord.assignment_id == DutyAssignment.id
    ).join(
        DutySection, DutyAssignment.section_id == DutySection.id
    ).outerjoin(
        Teacher, DutyAssignment.teacher_id == Teacher.id
    ).filter(
        AttendanceRecord.date == today,
        DutySection.plan_id == plan.id
    ).all()
    
    results = {
        'absent_warnings': [],
        'absent_deductions': [],
        'late_warnings': [],
        'late_deductions': [],
        'errors': [],
        'skipped': []
    }
    
    # Process each attendance record
    for record, assignment, section, teacher in attendance_records:
        if not teacher:
            results['skipped'].append({
                'name': assignment.placeholder_name or 'Unknown',
                'reason': 'No teacher linked to assignment'
            })
            continue
        
        if record.status == 'absent':
            # Increment warning count
            teacher.warnings += 1

            if teacher.warnings >= 3:
                # 3rd absence - send 1 FULL DAY deduction email and reset
                email_result = send_absence_deduction_email(
                    teacher_email=teacher.email,
                    teacher_name=teacher.full_name,
                    attendance_date=today,
                    day=today_day_name,
                    section=section.name,
                    task_place=assignment.place_task,
                    deduction_type="full_day",
                    warning_number=3
                )
                teacher.warnings = 0  # Reset after deduction

                results['absent_deductions'].append({
                    'name': teacher.full_name,
                    'email': teacher.email,
                    'deduction': '1 day',
                    'result': email_result
                })
            elif teacher.warnings == 2:
                # 2nd absence - send 1/2 DAY deduction email
                email_result = send_absence_deduction_email(
                    teacher_email=teacher.email,
                    teacher_name=teacher.full_name,
                    attendance_date=today,
                    day=today_day_name,
                    section=section.name,
                    task_place=assignment.place_task,
                    deduction_type="half_day",
                    warning_number=2
                )

                results['absent_deductions'].append({
                    'name': teacher.full_name,
                    'email': teacher.email,
                    'deduction': '1/2 day',
                    'result': email_result
                })
            else:
                # 1st absence - send warning email
                email_result = send_absence_warning_email(
                    teacher_email=teacher.email,
                    teacher_name=teacher.full_name,
                    warning_number=teacher.warnings,
                    attendance_date=today,
                    day=today_day_name,
                    section=section.name,
                    task_place=assignment.place_task
                )

                results['absent_warnings'].append({
                    'name': teacher.full_name,
                    'email': teacher.email,
                    'warning_number': teacher.warnings,
                    'result': email_result
                })
        
        elif record.status == 'late':
            # Increment late count
            teacher.late_count += 1

            if teacher.late_count >= 3:
                # 3rd late - send 1/2 DAY deduction email and reset
                email_result = send_late_deduction_email(
                    teacher_email=teacher.email,
                    teacher_name=teacher.full_name,
                    attendance_date=today,
                    day=today_day_name,
                    section=section.name,
                    task_place=assignment.place_task,
                    deduction_type="half_day",
                    late_number=3
                )
                teacher.late_count = 0  # Reset after deduction

                results['late_deductions'].append({
                    'name': teacher.full_name,
                    'email': teacher.email,
                    'deduction': '1/2 day',
                    'result': email_result
                })
            elif teacher.late_count == 2:
                # 2nd late - send 1/4 DAY deduction email
                email_result = send_late_deduction_email(
                    teacher_email=teacher.email,
                    teacher_name=teacher.full_name,
                    attendance_date=today,
                    day=today_day_name,
                    section=section.name,
                    task_place=assignment.place_task,
                    deduction_type="quarter_day",
                    late_number=2
                )

                results['late_deductions'].append({
                    'name': teacher.full_name,
                    'email': teacher.email,
                    'deduction': '1/4 day',
                    'result': email_result
                })
            else:
                # 1st late - send warning email
                email_result = send_late_warning_email(
                    teacher_email=teacher.email,
                    teacher_name=teacher.full_name,
                    late_number=teacher.late_count,
                    attendance_date=today,
                    day=today_day_name,
                    section=section.name,
                    task_place=assignment.place_task
                )

                results['late_warnings'].append({
                    'name': teacher.full_name,
                    'email': teacher.email,
                    'late_number': teacher.late_count,
                    'result': email_result
                })
    
    # Commit all changes
    db.session.commit()
    
    # Build summary
    summary = {
        'total_absent_warnings': len(results['absent_warnings']),
        'total_absent_deductions': len(results['absent_deductions']),
        'total_late_warnings': len(results['late_warnings']),
        'total_late_deductions': len(results['late_deductions']),
        'total_skipped': len(results['skipped']),
        'details': results
    }
    
    if request.headers.get('Accept') == 'application/json':
        return jsonify({'success': True, 'data': summary})
    
    # Flash summary messages
    if results['absent_warnings']:
        flash(f"Sent {len(results['absent_warnings'])} absence warning(s)", 'info')
    if results['absent_deductions']:
        flash(f"Sent {len(results['absent_deductions'])} absence deduction notice(s) (1/5 day)", 'warning')
    if results['late_warnings']:
        flash(f"Sent {len(results['late_warnings'])} late warning(s)", 'info')
    if results['late_deductions']:
        flash(f"Sent {len(results['late_deductions'])} late deduction notice(s) (half day)", 'warning')
    if results['skipped']:
        flash(f"Skipped {len(results['skipped'])} record(s) - no teacher linked", 'secondary')
    
    if not any([results['absent_warnings'], results['absent_deductions'], 
                results['late_warnings'], results['late_deductions']]):
        flash('No absent or late teachers found for today', 'info')
    
    return redirect(url_for('main.today_plan'))


@main_bp.route("/test-email")
def test_email():
    """
    Test route to verify email system is working.
    Sends test emails to ibrahim.fakhreyams@gmail.com
    """
    from datetime import date
    
    test_email_address = 'ibrahim.fakhreyams@gmail.com'
    today = date.today()
    today_day_name = today.strftime("%A")
    
    results = {
        'absence_warning': None,
        'absence_deduction': None,
        'late_warning': None,
        'late_deduction': None
    }
    
    # Test 1: Absence Warning Email
    results['absence_warning'] = send_absence_warning_email(
        teacher_email=test_email_address,
        teacher_name='Test Teacher',
        warning_number=1,
        attendance_date=today,
        day=today_day_name,
        section='Test Section',
        task_place='Test Location'
    )
    
    # Test 2: Absence Deduction Email (1/5 day)
    results['absence_deduction'] = send_absence_deduction_email(
        teacher_email=test_email_address,
        teacher_name='Test Teacher',
        attendance_date=today,
        day=today_day_name,
        section='Test Section',
        task_place='Test Location'
    )
    
    # Test 3: Late Warning Email
    results['late_warning'] = send_late_warning_email(
        teacher_email=test_email_address,
        teacher_name='Test Teacher',
        late_number=2,
        attendance_date=today,
        day=today_day_name,
        section='Test Section',
        task_place='Test Location'
    )
    
    # Test 4: Late Deduction Email (half day)
    results['late_deduction'] = send_late_deduction_email(
        teacher_email=test_email_address,
        teacher_name='Test Teacher',
        attendance_date=today,
        day=today_day_name,
        section='Test Section',
        task_place='Test Location'
    )
    
    # Build response
    success_count = sum(1 for r in results.values() if r and r.get('success'))

    return jsonify({
        'success': success_count == 4,
        'message': f'Sent {success_count}/4 test emails to {test_email_address}',
        'results': results
    })


@main_bp.route("/system-test")
def system_test():
    """
    Simple test route to verify the email system is working.
    Sends a test email to all administrators.
    """
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    from datetime import datetime

    results = {
        'smtp_connection': False,
        'email_sent': False,
        'admins': ADMIN_CC_EMAILS,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'errors': []
    }

    try:
        # Test 1: SMTP Connection
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            results['smtp_connection'] = True

            # Test 2: Send test email to admins
            msg = MIMEMultipart('alternative')
            msg['From'] = EMAIL_ADDRESS
            msg['To'] = ', '.join(ADMIN_CC_EMAILS)
            msg['Subject'] = f"🔧 System Test - Dismissal Checker Working - {results['timestamp']}"

            html_body = f"""
            <html>
            <body style="font-family: Arial, sans-serif; padding: 20px;">
                <h2 style="color: #28a745;">✅ System Test Successful</h2>
                <p>This is a test email to confirm the Dismissal Checker email system is working correctly.</p>
                <hr>
                <table style="border-collapse: collapse; width: 100%; max-width: 500px;">
                    <tr>
                        <td style="padding: 8px; border: 1px solid #ddd;"><strong>Timestamp:</strong></td>
                        <td style="padding: 8px; border: 1px solid #ddd;">{results['timestamp']}</td>
                    </tr>
                    <tr>
                        <td style="padding: 8px; border: 1px solid #ddd;"><strong>SMTP Server:</strong></td>
                        <td style="padding: 8px; border: 1px solid #ddd;">{SMTP_SERVER}:{SMTP_PORT}</td>
                    </tr>
                    <tr>
                        <td style="padding: 8px; border: 1px solid #ddd;"><strong>From:</strong></td>
                        <td style="padding: 8px; border: 1px solid #ddd;">{EMAIL_ADDRESS}</td>
                    </tr>
                    <tr>
                        <td style="padding: 8px; border: 1px solid #ddd;"><strong>Status:</strong></td>
                        <td style="padding: 8px; border: 1px solid #ddd; color: green;">All Systems Operational</td>
                    </tr>
                </table>
                <hr>
                <p style="color: #666; font-size: 12px;">This is an automated test email from the Dismissal Checker system.</p>
            </body>
            </html>
            """

            html_part = MIMEText(html_body, 'html')
            msg.attach(html_part)

            server.sendmail(EMAIL_ADDRESS, ADMIN_CC_EMAILS, msg.as_string())
            results['email_sent'] = True

    except smtplib.SMTPAuthenticationError as e:
        results['errors'].append(f"SMTP Authentication Error: {str(e)}")
    except smtplib.SMTPException as e:
        results['errors'].append(f"SMTP Error: {str(e)}")
    except Exception as e:
        results['errors'].append(f"General Error: {str(e)}")

    # Build response
    success = results['smtp_connection'] and results['email_sent']

    return jsonify({
        'success': success,
        'message': 'System test completed successfully! Email sent to all admins.' if success else 'System test failed.',
        'details': results
    })


@main_bp.route("/reset-all-warnings")
def reset_all_warnings():
    """
    Reset all teachers' warnings and late counts to 0.
    """
    try:
        count = Teacher.query.update({Teacher.warnings: 0, Teacher.late_count: 0})
        db.session.commit()
        
        if request.headers.get('Accept') == 'application/json':
            return jsonify({'success': True, 'message': f'Reset warnings for {count} teachers'})
        
        flash(f'Successfully reset warnings for {count} teachers', 'success')
        return redirect(url_for('main.teachers'))
    except Exception as e:
        db.session.rollback()
        if request.headers.get('Accept') == 'application/json':
            return jsonify({'success': False, 'message': str(e)}), 500
        flash(f'Error resetting warnings: {str(e)}', 'error')
        return redirect(url_for('main.teachers'))


@main_bp.route("/plan/today", methods=["GET", "POST"])
def today_plan():
    today = date.today()
    return plan_for_day(today.strftime("%A"), today)


@main_bp.route("/plan/<day>", methods=["GET", "POST"])
def plan_for_day(day: str, target_date: date | None = None):
    day = day.capitalize()
    if target_date is None:
        date_param = request.args.get("date") or request.form.get("date")
        if date_param:
            try:
                target_date = datetime.strptime(date_param, "%Y-%m-%d").date()
            except ValueError:
                flash("Invalid date format. Please use YYYY-MM-DD.", "warning")
                target_date = date.today()
        else:
            target_date = date.today()

    all_plans = DutyPlan.query.order_by(DutyPlan.day_of_week).all()
    plan = DutyPlan.query.filter_by(day_of_week=day).first()
    if plan is None:
        flash(f"No dismissal plan found for {day}.", "warning")
        return render_template("plan.html", plan=None, day=day, target_date=target_date, all_plans=all_plans)

    if request.method == "POST":
        # Check if this is an offline sync request
        is_offline_sync = request.headers.get('X-Offline-Sync') == 'true'
        original_timestamp = request.headers.get('X-Original-Timestamp')

        if is_offline_sync:
            # Handle offline sync with conflict resolution
            result = handle_offline_sync_submission(plan, target_date, original_timestamp)
            if result['success']:
                return {"status": "success", "message": result['message']}, 200
            else:
                return {"status": "error", "message": result['message']}, 400
        else:
            # Handle normal submission
            handle_attendance_submission(plan, target_date)
            return redirect(url_for("main.plan_for_day", day=day, date=target_date.isoformat()))

    attendance_map = load_attendance_map(plan, target_date)
    return render_template(
        "plan.html",
        plan=plan,
        attendance_map=attendance_map,
        day=day,
        target_date=target_date,
        all_plans=all_plans,
    )


@main_bp.route("/teachers")
def teachers():
    teacher_list = Teacher.query.order_by(Teacher.full_name).all()
    return render_template("teachers.html", teachers=teacher_list)


@main_bp.route("/import/teachers", methods=["GET", "POST"])
def import_teachers():
    if request.method == "POST":
        file = request.files.get("file")
        if not file or file.filename == "":
            flash("Please choose an .xls file to upload.", "danger")
            return redirect(request.url)

        filename = secure_filename(file.filename)
        if not filename.lower().endswith(".xls"):
            flash("Only .xls files are supported.", "danger")
            return redirect(request.url)

        upload_dir = Path(current_app.instance_path)
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_path = upload_dir / filename
        file.save(file_path)

        try:
            imported_count = import_teachers_from_xls(file_path)
            flash(f"Successfully imported {imported_count} teachers.", "success")
        except Exception as exc:  # pragma: no cover - user feedback
            current_app.logger.exception("Failed to import teachers")
            flash(f"Failed to import teachers: {exc}", "danger")
        finally:
            if file_path.exists():
                file_path.unlink()

        return redirect(url_for("main.teachers"))

    return render_template("import_teachers.html")


@main_bp.route("/send-system-announcement", methods=["GET", "POST"])
def send_system_announcement():
    """Send an email to all teachers about the new warning/deduction system."""
    teachers = Teacher.query.filter(Teacher.email.isnot(None)).all()

    if request.method == "POST":
        results = {'sent': 0, 'failed': 0, 'messages': []}

        for teacher in teachers:
            if teacher.email:
                result = send_new_system_announcement(teacher.email, teacher.full_name)
                if result['success']:
                    results['sent'] += 1
                else:
                    results['failed'] += 1
                results['messages'].append(result['message'])

        flash(f"Announcement sent to {results['sent']} teachers. {results['failed']} failed.",
              "success" if results['failed'] == 0 else "warning")
        return redirect(url_for("main.send_system_announcement"))

    return render_template("send_announcement.html", teachers=teachers)


def import_teachers_from_xls(path: Path) -> int:
    workbook = open_workbook(path, encoding_override="utf-8")
    sheet = workbook.sheet_by_index(0)

    headers = [str(sheet.cell_value(0, col)).strip().lower() for col in range(sheet.ncols)]
    expected_columns = {"num", "full name", "mobile", "email"}
    if not expected_columns.issubset(set(headers)):
        raise ValueError("The uploaded file does not contain the required columns: NUM, Full name, Mobile, Email")

    name_idx = headers.index("full name")
    mobile_idx = headers.index("mobile")
    email_idx = headers.index("email")

    count = 0
    for row_idx in range(1, sheet.nrows):
        full_name = to_string(sheet.cell(row_idx, name_idx))
        mobile = to_string(sheet.cell(row_idx, mobile_idx))
        email = to_string(sheet.cell(row_idx, email_idx)).lower()
        if not full_name or not email:
            continue

        teacher = Teacher.query.filter_by(email=email).first()
        if teacher:
            teacher.full_name = normalize_name(full_name)
            teacher.mobile = mobile
        else:
            teacher = Teacher(full_name=normalize_name(full_name), mobile=mobile, email=email)
            db.session.add(teacher)
        count += 1

    db.session.commit()
    return count


def normalize_name(name: str) -> str:
    return " ".join(part for part in name.split())


def to_string(cell) -> str:
    value = cell.value
    if cell.ctype == XL_CELL_NUMBER:
        if float(value).is_integer():
            return str(int(value))
        return str(value)
    return str(value).strip()


def handle_offline_sync_submission(plan: DutyPlan, target_date: date, original_timestamp: str | None) -> dict:
    """
    Handle offline sync submission for attendance data.

    Args:
        plan: The duty plan
        target_date: The date for the attendance
        original_timestamp: When the offline submission was originally made

    Returns:
        dict: Result with success status and message
    """
    try:
        # Check for conflicts - if there are newer records, we need to handle them
        conflicts = []
        submission_count = 0
        notification_teachers = []

        # Use no_autoflush to prevent premature flush during iteration
        with db.session.no_autoflush:
            for assignment in plan_assignments(plan):
                form_key = f"assignment-{assignment.id}"
                notes_key = f"notes-{assignment.id}"
                status = request.form.get(form_key)
                notes = request.form.get(notes_key, "").strip()

                if status:  # Only process if status is provided
                    # Get existing record
                    existing_record = AttendanceRecord.query.filter_by(
                        assignment_id=assignment.id,
                        date=target_date
                    ).first()

                    previous_status = existing_record.status if existing_record else None

                    if existing_record:
                        # Update existing record instead of delete+add
                        existing_record.status = status
                        existing_record.notes = notes or None
                    else:
                        # Create new record only if none exists
                        new_record = AttendanceRecord(
                            assignment_id=assignment.id,
                            date=target_date,
                            status=status,
                            notes=notes or None
                        )
                        db.session.add(new_record)

                    # Adjust warnings based on status change
                    adjust_warnings(assignment, previous_status, status)

                    submission_count += 1

                    # Track teachers for notifications
                    if assignment.teacher and status in ['absent', 'late']:
                        notification_teachers.append({
                            'teacher': assignment.teacher,
                            'status': status,
                            'assignment': assignment,
                            'date': target_date
                        })

        # Commit all changes
        db.session.commit()

        # Send email notifications for absent/late teachers
        if notification_teachers:
            try:
                send_bulk_attendance_notifications(notification_teachers)
            except Exception as e:
                # Log error but don't fail the sync
                pass

        return {
            'success': True,
            'message': f'Successfully synced {submission_count} attendance records',
            'conflicts': conflicts
        }

    except Exception as e:
        db.session.rollback()
        return {
            'success': False,
            'message': f'Error during sync: {str(e)}'
        }


def handle_attendance_submission(plan: DutyPlan, target_date: date) -> None:
    """Handle attendance form submission - only saves data without sending emails"""
    submission_count = 0

    # Use no_autoflush to prevent premature flush during iteration
    with db.session.no_autoflush:
        for assignment in plan_assignments(plan):
            form_key = f"assignment-{assignment.id}"
            notes_key = f"notes-{assignment.id}"
            status = request.form.get(form_key)
            notes = request.form.get(notes_key, "").strip()

            if status:  # Only process if status is provided
                # Get existing record
                existing_record = AttendanceRecord.query.filter_by(
                    assignment_id=assignment.id,
                    date=target_date
                ).first()

                previous_status = existing_record.status if existing_record else None

                if existing_record:
                    # Update existing record instead of delete+add
                    existing_record.status = status
                    existing_record.notes = notes or None
                else:
                    # Create new record only if none exists
                    new_record = AttendanceRecord(
                        assignment_id=assignment.id,
                        date=target_date,
                        status=status,
                        notes=notes or None
                    )
                    db.session.add(new_record)

                # Adjust warnings based on status change
                adjust_warnings(assignment, previous_status, status)

                submission_count += 1

    # Commit all changes
    db.session.commit()

    # Note: Email notifications are now handled by the scheduled task at 4:00 PM


def load_attendance_map(plan: DutyPlan, target_date: date) -> dict[int, AttendanceRecord | None]:
    records = (
        AttendanceRecord.query.join(DutyAssignment)
        .filter(DutyAssignment.section.has(plan_id=plan.id), AttendanceRecord.date == target_date)
        .all()
    )
    return {record.assignment_id: record for record in records}


def adjust_warnings(assignment: DutyAssignment, previous_status: str | None, new_status: str | None) -> None:
    """Adjust teacher warnings based on attendance status changes"""
    teacher = assignment.teacher
    if not teacher:
        return

    # Decrease warnings if changing from absent/late to present
    if previous_status in ['absent', 'late'] and new_status == 'present':
        teacher.warnings = max(0, teacher.warnings - 1)

    # Increase warnings if changing to absent/late
    elif new_status in ['absent', 'late'] and previous_status != new_status:
        teacher.warnings += 1


def plan_assignments(plan: DutyPlan):
    for section in plan.sections:
        for assignment in section.assignments:
            yield assignment


@main_bp.route("/upload/plan", methods=["GET", "POST"])
def upload_plan():
    """Upload a Word document containing the dismissal plan."""
    if request.method == "POST":
        if "plan_file" not in request.files:
            flash("No file selected", "error")
            return redirect(request.url)

        file = request.files["plan_file"]
        if file.filename == "":
            flash("No file selected", "error")
            return redirect(request.url)

        if file and file.filename and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            # Add timestamp to avoid filename conflicts
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            name, ext = os.path.splitext(filename)
            filename = f"{name}_{timestamp}{ext}"

            # Save to resources directory
            resources_dir = Path(current_app.root_path).parent / "resources"
            resources_dir.mkdir(exist_ok=True)
            file_path = resources_dir / filename

            try:
                file.save(file_path)

                # Process the Word document
                plan_data = process_word_document(file_path)

                flash(f"Plan uploaded successfully as {filename}. Extracted {len(plan_data)} items.", "success")
                return redirect(url_for("main.resources"))

            except Exception as e:
                flash(f"Error processing file: {str(e)}", "error")
                return redirect(request.url)
        else:
            flash("Invalid file type. Please upload a Word document (.docx)", "error")
            return redirect(request.url)

    return render_template("upload_plan.html")


@main_bp.route("/resources")
def resources():
    """Display uploaded resources for teachers."""
    resources_dir = Path(current_app.root_path).parent / "resources"
    resources_dir.mkdir(exist_ok=True)

    files = []
    if resources_dir.exists():
        for file_path in resources_dir.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in ['.docx', '.doc']:
                files.append({
                    'name': file_path.name,
                    'size': file_path.stat().st_size,
                    'modified': datetime.fromtimestamp(file_path.stat().st_mtime)
                })

    return render_template("resources.html", files=files)


@main_bp.route("/download/<filename>")
def download_file(filename):
    """Download a resource file."""
    resources_dir = Path(current_app.root_path).parent / "resources"
    return send_from_directory(resources_dir, filename, as_attachment=True)


@main_bp.route('/assign_roles', methods=['GET', 'POST'])
def assign_roles():
    """Assign roles to teachers based on parsed plan data"""
    from .models import Teacher, DutyAssignment, DutySection, DutyPlan
    import json

    if request.method == 'POST':
        # Load parsed assignments
        assignments_file = os.path.join(current_app.root_path, '..', 'data', 'parsed_assignments.json')

        if not os.path.exists(assignments_file):
            flash('No parsed assignments found. Please upload and parse a plan first.', 'error')
            return redirect(url_for('main.assign_roles'))

        with open(assignments_file, 'r', encoding='utf-8') as f:
            assignments_data = json.load(f)

        # Clear existing dismissal plans
        existing_plans = DutyPlan.query.filter(
            (DutyPlan.name.like('%Dismissal%')) | (DutyPlan.is_daily_plan == True)
        ).all()
        for plan in existing_plans:
            db.session.delete(plan)
        db.session.commit()

        # Create duty plans for each day with daily assignments
        for day, day_data in assignments_data.get('daily_assignments', {}).items():
            dismissal_plan = DutyPlan(
                name=f'{day} Dismissal',
                day_of_week=day,
                supervisor=day_data.get('supervisor'),
                team=day_data.get('team'),
                is_daily_plan=True
            )
            db.session.add(dismissal_plan)
            db.session.flush()  # Get the ID

            # Group teachers by role for this day
            role_groups = {}
            for teacher_data in day_data.get('teachers', []):
                role = teacher_data.get('role', 'General Duty')
                if role not in role_groups:
                    role_groups[role] = []
                role_groups[role].append(teacher_data)

            # Create sections and assignments for this day's plan
            section_order = 0
            for role, teachers in role_groups.items():
                section_order += 1
                section = DutySection(
                    plan_id=dismissal_plan.id,
                    name=role,
                    order=section_order
                )
                db.session.add(section)
                db.session.flush()  # Get the section ID

                # Create assignments for each teacher in this role
                assignment_order = 0
                for teacher_data in teachers:
                    assignment_order += 1

                    # Try to find existing teacher by name
                    teacher = Teacher.query.filter_by(full_name=teacher_data['name']).first()

                    assignment = DutyAssignment(
                        section_id=section.id,
                        teacher_id=teacher.id if teacher else None,
                        placeholder_name=teacher_data['name'] if not teacher else None,
                        order=assignment_order,
                        place_task=teacher_data.get('role', role)
                    )
                    db.session.add(assignment)

        # Create a general dismissal plan for all teacher assignments (fallback)
        if assignments_data.get('teachers'):
            general_plan = DutyPlan(
                name='General Dismissal',
                day_of_week=None,
                supervisor='Dismissal Coordinator',
                team='All Staff',
                is_daily_plan=False
            )
            db.session.add(general_plan)
            db.session.flush()

            # Group teachers by role
            role_groups = {}
            for teacher_data in assignments_data['teachers']:
                role = teacher_data['role']
                if role not in role_groups:
                    role_groups[role] = []
                role_groups[role].append(teacher_data)

            # Create sections and assignments for general plan
            section_order = 0
            for role, teachers in role_groups.items():
                section_order += 1
                section = DutySection(
                    plan_id=general_plan.id,
                    name=role,
                    order=section_order
                )
                db.session.add(section)
                db.session.flush()  # Get the section ID

                # Create assignments for each teacher in this role
                assignment_order = 0
                for teacher_data in teachers:
                    assignment_order += 1

                    # Try to find existing teacher by name
                    teacher = Teacher.query.filter_by(full_name=teacher_data['name']).first()

                    assignment = DutyAssignment(
                        section_id=section.id,
                        teacher_id=teacher.id if teacher else None,
                        placeholder_name=teacher_data['name'] if not teacher else None,
                        order=assignment_order,
                        place_task=role
                    )
                    db.session.add(assignment)

        db.session.commit()

        total_daily_assignments = len(assignments_data.get('daily_assignments', {}))
        total_teachers_assigned = sum(len(day_data.get('teachers', [])) for day_data in assignments_data.get('daily_assignments', {}).values())
        total_general_teachers = len(assignments_data.get('teachers', []))

        flash(f'Successfully created {total_daily_assignments} daily plans with {total_teachers_assigned} teacher assignments! Also created general plan with {total_general_teachers} teachers.', 'success')
        return redirect(url_for('main.view_assignments'))

    # GET request - show assignment form
    assignments_file = os.path.join(current_app.root_path, '..', 'data', 'parsed_assignments.json')
    assignments_data = None

    if os.path.exists(assignments_file):
        with open(assignments_file, 'r', encoding='utf-8') as f:
            assignments_data = json.load(f)

    return render_template('assign_roles.html', assignments_data=assignments_data)


@main_bp.route('/view_assignments')
def view_assignments():
    """View all duty assignments"""
    try:
        # Get all duty plans
        plans = DutyPlan.query.all()

        assignments_data = []
        for plan in plans:
            for section in plan.sections:
                for assignment in section.assignments:
                    assignments_data.append({
                        'plan_name': plan.name,
                        'day_of_week': plan.day_of_week,
                        'supervisor': plan.supervisor,
                        'section_name': section.name,
                        'teacher_name': assignment.display_name,
                        'place_task': assignment.place_task
                    })

        return render_template('view_assignments.html', assignments=assignments_data)
    except Exception as e:
        flash(f'Error loading assignments: {str(e)}', 'danger')
        return redirect(url_for('main.index'))


@main_bp.route('/reports')
def attendance_reports():
    """View attendance reports"""
    try:
        # Get filter parameters
        date_filter = request.args.get('date')
        teacher_filter = request.args.get('teacher')
        status_filter = request.args.get('status')

        # Base query with joins
        query = db.session.query(AttendanceRecord).join(
            DutyAssignment, AttendanceRecord.assignment_id == DutyAssignment.id
        ).join(
            DutySection, DutyAssignment.section_id == DutySection.id
        ).join(
            DutyPlan, DutySection.plan_id == DutyPlan.id
        ).outerjoin(
            Teacher, DutyAssignment.teacher_id == Teacher.id
        )

        # Apply filters
        if date_filter:
            try:
                filter_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
                query = query.filter(AttendanceRecord.date == filter_date)
            except ValueError:
                flash('Invalid date format', 'warning')

        if teacher_filter:
            query = query.filter(
                db.or_(
                    Teacher.full_name.ilike(f'%{teacher_filter}%'),
                    DutyAssignment.placeholder_name.ilike(f'%{teacher_filter}%')
                )
            )

        if status_filter and status_filter in ['present', 'absent']:
            query = query.filter(AttendanceRecord.status == status_filter)

        # Order by date descending, then by plan name
        records = query.order_by(
            AttendanceRecord.date.desc(),
            DutyPlan.name,
            DutySection.order,
            DutyAssignment.order
        ).all()

        # Get unique dates and teachers for filter dropdowns
        all_dates = db.session.query(AttendanceRecord.date).distinct().order_by(AttendanceRecord.date.desc()).all()
        unique_dates = [d[0] for d in all_dates]

        # Get all teachers (both assigned and placeholder names)
        assigned_teachers = db.session.query(Teacher.full_name).distinct().order_by(Teacher.full_name).all()
        placeholder_teachers = db.session.query(DutyAssignment.placeholder_name).filter(
            DutyAssignment.placeholder_name.isnot(None)
        ).distinct().order_by(DutyAssignment.placeholder_name).all()

        all_teachers = sorted(set([t[0] for t in assigned_teachers] + [t[0] for t in placeholder_teachers if t[0]]))

        return render_template('reports.html',
                             records=records,
                             unique_dates=unique_dates,
                             all_teachers=all_teachers,
                             current_filters={
                                 'date': date_filter,
                                 'teacher': teacher_filter,
                                 'status': status_filter
                             })
    except Exception as e:
        flash(f'Error loading reports: {str(e)}', 'danger')
        return redirect(url_for('main.index'))


def allowed_file(filename):
    """Check if the uploaded file has an allowed extension."""
    ALLOWED_EXTENSIONS = {'docx', 'doc'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def process_word_document(file_path):
    """Process a Word document and extract plan data."""
    try:
        doc = Document(file_path)
        plan_data = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                plan_data.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    plan_data.append(" | ".join(row_data))

        return plan_data

    except Exception as e:
        raise Exception(f"Failed to process Word document: {str(e)}")


@main_bp.route('/api/edit-plan', methods=['POST'])
def edit_daily_plan():
    """
    API route to edit a daily plan by specifying the day and teacher name.

    Expected JSON payload:
    {
        "day": "Monday",
        "teacher_name": "John Doe",
        "new_teacher_name": "Jane Smith",  # optional, for updating teacher name
        "teacher_email": "jane.smith@example.com"
    }
    """
    try:
        # Validate request content type
        if not request.is_json:
            return jsonify({
                "status": "error",
                "message": "Content-Type must be application/json"
            }), 400

        data = request.get_json()

        # Validate required parameters
        if not data:
            return jsonify({
                "status": "error",
                "message": "No JSON data provided"
            }), 400

        day = data.get('day')
        teacher_name = data.get('teacher_name')
        new_teacher_name = data.get('new_teacher_name', teacher_name)
        teacher_email = data.get('teacher_email')

        # Validate required fields
        if not day:
            return jsonify({
                "status": "error",
                "message": "Day parameter is required"
            }), 400

        if not teacher_name:
            return jsonify({
                "status": "error",
                "message": "Teacher name parameter is required"
            }), 400

        if not teacher_email:
            return jsonify({
                "status": "error",
                "message": "Teacher email parameter is required"
            }), 400

        # Validate day format
        day = day.capitalize()
        valid_days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
        if day not in valid_days:
            return jsonify({
                "status": "error",
                "message": f"Invalid day. Must be one of: {', '.join(valid_days)}"
            }), 400

        # Validate email format (basic validation)
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, teacher_email):
            return jsonify({
                "status": "error",
                "message": "Invalid email format"
            }), 400

        # Find the duty plan for the specified day
        plan = DutyPlan.query.filter_by(day_of_week=day).first()
        if not plan:
            return jsonify({
                "status": "error",
                "message": f"No duty plan found for {day}"
            }), 404

        # Find the teacher in the plan
        assignment = None
        for section in plan.sections:
            for assign in section.assignments:
                # Check both assigned teacher and placeholder name
                if (assign.teacher and assign.teacher.full_name == teacher_name) or \
                   (assign.placeholder_name == teacher_name):
                    assignment = assign
                    break
            if assignment:
                break

        if not assignment:
            return jsonify({
                "status": "error",
                "message": f"Teacher '{teacher_name}' not found in {day} plan"
            }), 404

        # Find the existing teacher by email
        existing_teacher = Teacher.query.filter_by(email=teacher_email).first()

        if not existing_teacher:
            return jsonify({
                "status": "error",
                "message": f"Teacher with email '{teacher_email}' not found in the system"
            }), 404

        # Simply assign the existing teacher to the assignment
        # Don't update teacher information, just link the assignment to the teacher
        assignment.teacher_id = existing_teacher.id
        assignment.placeholder_name = None

        # Commit changes
        db.session.commit()

        return jsonify({
            "status": "success",
            "message": f"Successfully assigned teacher to {day} plan",
            "data": {
                "day": day,
                "teacher_name": existing_teacher.full_name,
                "teacher_email": existing_teacher.email,
                "section": assignment.section.name,
                "place_task": assignment.place_task
            }
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({
            "status": "error",
            "message": f"An error occurred: {str(e)}"
        }), 500


@main_bp.route('/api/manage-warnings', methods=['POST'])
def manage_warnings():
    """
    Manage teacher warnings - reset, reduce by one, or set to a specific number

    Expected JSON payload:
    {
        "teacher_name": "Teacher Name",
        "action": "reset|reduce|set",
        "warnings": 5  // Only required for "set" action
    }
    """
    try:
        data = request.get_json()

        if not data:
            return jsonify({
                'success': False,
                'message': 'No data provided'
            }), 400

        teacher_name = data.get('teacher_name', '').strip()
        action = data.get('action', '').strip().lower()

        if not teacher_name:
            return jsonify({
                'success': False,
                'message': 'Teacher name is required'
            }), 400

        if action not in ['reset', 'reduce', 'set']:
            return jsonify({
                'success': False,
                'message': 'Action must be "reset", "reduce", or "set"'
            }), 400

        # Find teacher by name (case-insensitive search)
        teacher = Teacher.query.filter(
            Teacher.full_name.ilike(f'%{teacher_name}%')
        ).first()

        if not teacher:
            return jsonify({
                'success': False,
                'message': f'Teacher "{teacher_name}" not found'
            }), 404

        old_warnings = teacher.warnings

        # Perform the requested action
        if action == 'reset':
            teacher.warnings = 0
        elif action == 'reduce':
            teacher.warnings = max(0, teacher.warnings - 1)
        elif action == 'set':
            warnings_value = data.get('warnings')
            if warnings_value is None:
                return jsonify({
                    'success': False,
                    'message': 'Warnings value is required for "set" action'
                }), 400

            try:
                warnings_value = int(warnings_value)
                if warnings_value < 0:
                    return jsonify({
                        'success': False,
                        'message': 'Warnings value must be non-negative'
                    }), 400
                teacher.warnings = warnings_value
            except (ValueError, TypeError):
                return jsonify({
                    'success': False,
                    'message': 'Warnings value must be a valid number'
                }), 400

        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Successfully updated warnings for {teacher.full_name}',
            'teacher': {
                'name': teacher.full_name,
                'email': teacher.email,
                'old_warnings': old_warnings,
                'new_warnings': teacher.warnings
            }
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': f'Error managing warnings: {str(e)}'
        }), 500


@main_bp.route('/edit-plan')
def edit_plan_form():
    """Serve the edit plan form page."""
    return render_template('edit_plan.html')


@main_bp.route('/manage-warnings')
def manage_warnings_page():
    """Serve the manage warnings HTML page"""
    return render_template('manage_warnings.html')


@main_bp.route('/api/search-teachers', methods=['GET'])
def search_teachers():
    """Search for teachers by name or email for autocomplete functionality."""
    query = request.args.get('q', '').strip()
    search_type = request.args.get('type', 'name')  # 'name' or 'email'

    if not query or len(query) < 2:
        return jsonify([])

    try:
        if search_type == 'email':
            # Search by email
            teachers = Teacher.query.filter(
                Teacher.email.like(f'%{query}%')
            ).limit(10).all()

            results = [{'email': teacher.email, 'name': teacher.full_name} for teacher in teachers]
        else:
            # Search by name (default)
            teachers = Teacher.query.filter(
                Teacher.full_name.like(f'%{query}%')
            ).limit(10).all()

            results = [{'name': teacher.full_name, 'email': teacher.email} for teacher in teachers]

            # Also search for placeholder names in duty assignments
            from sqlalchemy import text
            placeholder_query = text("""
                SELECT DISTINCT placeholder_name
                FROM duty_assignments
                WHERE placeholder_name IS NOT NULL
                AND LOWER(placeholder_name) LIKE LOWER(:query)
                LIMIT 5
            """)

            placeholder_results = db.session.execute(placeholder_query, {'query': f'%{query}%'}).fetchall()

            # Add placeholder results
            for row in placeholder_results:
                placeholder_name = row[0]
                if placeholder_name and placeholder_name not in [r['name'] for r in results]:
                    results.append({'name': placeholder_name, 'email': '', 'is_placeholder': True})

        return jsonify(results)

    except Exception as e:
        flash(f'Error processing absent teacher notifications: {str(e)}', 'error')
        return redirect(url_for('main.index'))


@main_bp.route('/scheduled-tasks', methods=['POST'])
def run_scheduled_tasks():
    """
    Route for scheduled tasks that runs at 4:00 PM Cairo time.
    Handles both daily report and individual notifications.
    """
    try:
        # Get current date in Cairo timezone
        cairo_tz = pytz.timezone('Africa/Cairo')
        cairo_now = datetime.now(cairo_tz)
        today = cairo_now.date()
        day_name = today.strftime("%A")

        # Get today's duty plan
        plan = DutyPlan.query.filter_by(day_of_week=day_name).first()

        if not plan:
            return jsonify({
                'success': False,
                'message': f'No duty plan found for {day_name}',
                'date': today.isoformat(),
                'cairo_time': cairo_now.strftime('%Y-%m-%d %H:%M:%S %Z')
            }), 404

        # Get absent teachers for today
        absent_records = db.session.query(
            AttendanceRecord, Teacher, DutyAssignment, DutySection
        ).join(
            DutyAssignment, AttendanceRecord.assignment_id == DutyAssignment.id
        ).join(
            Teacher, DutyAssignment.teacher_id == Teacher.id
        ).join(
            DutySection, DutyAssignment.section_id == DutySection.id
        ).filter(
            AttendanceRecord.date == today,
            AttendanceRecord.status == 'absent',
            DutySection.plan_id == plan.id
        ).all()

        absent_teachers = []
        for record, teacher, assignment, section in absent_records:
            absent_teachers.append({
                'name': teacher.full_name,
                'email': teacher.email,
                'section': section.name,
                'task_place': assignment.place_task,
                'supervisor': plan.supervisor,
                'warnings': teacher.warnings,
                'notes': record.notes
            })

        # Task 1: Send daily report to specified emails
        daily_report_success = False
        daily_report_message = ""

        if absent_teachers:
            try:
                import smtplib
                from email.mime.text import MIMEText
                from email.mime.multipart import MIMEMultipart

                # Create email content
                email_content = f"""
                <h2>Daily Absent Teachers Report - {today.strftime('%Y-%m-%d')}</h2>
                <p>Total absent teachers: {len(absent_teachers)}</p>
                <table border="1" style="border-collapse: collapse; width: 100%;">
                    <tr style="background-color: #f2f2f2;">
                        <th style="padding: 8px;">Teacher Name</th>
                        <th style="padding: 8px;">Section</th>
                        <th style="padding: 8px;">Task/Place</th>
                        <th style="padding: 8px;">Supervisor</th>
                        <th style="padding: 8px;">Warnings</th>
                    </tr>
                """

                for teacher in absent_teachers:
                    email_content += f"""
                    <tr>
                        <td style="padding: 8px;">{teacher['name']}</td>
                        <td style="padding: 8px;">{teacher['section']}</td>
                        <td style="padding: 8px;">{teacher['task_place'] or 'N/A'}</td>
                        <td style="padding: 8px;">{teacher['supervisor'] or 'N/A'}</td>
                        <td style="padding: 8px;">{teacher['warnings'] or 0}</td>
                    </tr>
                    """

                email_content += "</table>"

                # Send email using SMTP
                recipients = ["islam.qamar@ams-benha.com", "ibrahimfakhreyams@gmail.com","yasser.alaraby@ams-benha.com","hadeer.tawfik@ams-benha.com","mohamad.mosalam@ams-benha.com"]

                msg = MIMEMultipart('alternative')
                msg['From'] = "amsprog2022@gmail.com"
                msg['To'] = ", ".join(recipients)
                msg['Subject'] = f"Daily Absent Teachers Report - {today.strftime('%Y-%m-%d')}"

                html_part = MIMEText(email_content, 'html')
                msg.attach(html_part)

                with smtplib.SMTP("smtp.gmail.com", 587) as server:
                    server.starttls()
                    server.login("amsprog2022@gmail.com", "xfleslznraphvqgc")
                    server.send_message(msg)

                daily_report_success = True
                daily_report_message = f'Daily report sent successfully to {", ".join(recipients)} with {len(absent_teachers)} absent teachers'

            except Exception as email_error:
                daily_report_message = f'Failed to send daily report: {str(email_error)}'
        else:
            daily_report_success = True
            daily_report_message = 'No absent teachers today - no daily report sent'

        # Task 2: Send individual notifications to absent teachers
        individual_notifications_success = False
        individual_notifications_message = ""

        if absent_records:
            sent_count = 0
            skipped_count = 0
            failed_count = 0

            for record, teacher, assignment, section in absent_records:
                try:
                    result = send_attendance_notification(
                        teacher_email=teacher.email,
                        teacher_name=teacher.full_name,
                        status='absent',
                        day=day_name,
                        attendance_date=today,
                        section=section.name,
                        task_place=assignment.place_task,
                        supervisor=plan.supervisor,
                        warnings=teacher.warnings,
                        notes=record.notes
                    )

                    if result['success']:
                        if result.get('skipped', False):
                            skipped_count += 1
                        else:
                            sent_count += 1
                    else:
                        failed_count += 1

                except Exception as e:
                    failed_count += 1

            total_teachers = len(absent_records)
            individual_notifications_success = failed_count == 0
            individual_notifications_message = f'Individual notifications for {total_teachers} absent teachers: {sent_count} sent, {skipped_count} skipped, {failed_count} failed'
        else:
            individual_notifications_success = True
            individual_notifications_message = 'No absent teachers found for individual notifications'

        # Return comprehensive response
        return jsonify({
            'success': daily_report_success and individual_notifications_success,
            'date': today.isoformat(),
            'day': day_name,
            'cairo_time': cairo_now.strftime('%Y-%m-%d %H:%M:%S %Z'),
            'total_absent_teachers': len(absent_teachers),
            'daily_report': {
                'success': daily_report_success,
                'message': daily_report_message
            },
            'individual_notifications': {
                'success': individual_notifications_success,
                'message': individual_notifications_message
            }
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error in scheduled tasks: {str(e)}',
            'date': today.isoformat() if 'today' in locals() else None
        }), 500


@main_bp.route('/today-plan-and-absent', methods=['GET'])
def today_plan_and_absent():
    """
    Get today's plan, send email with absent teachers to ibrahimfakhreyams@gmail.com,
    and display HTML page with absent teachers.
    """
    try:
        from .email_service import send_attendance_notification

        today = date.today()
        today_day_name = today.strftime("%A")  # Get day name like "Monday", "Tuesday", etc.

        # Get today's plan
        plan = DutyPlan.query.filter_by(day_of_week=today_day_name).first()

        if not plan:
            flash('No plan found for today', 'warning')
            return render_template('absent_teachers_report.html',
                                 absent_teachers=[],
                                 plan_date=today,
                                 plan_exists=False)

        # Get absent teachers for today
        absent_records = db.session.query(
            AttendanceRecord, Teacher, DutyAssignment, DutySection
        ).join(
            DutyAssignment, AttendanceRecord.assignment_id == DutyAssignment.id
        ).join(
            Teacher, DutyAssignment.teacher_id == Teacher.id
        ).join(
            DutySection, DutyAssignment.section_id == DutySection.id
        ).filter(
            AttendanceRecord.date == today,
            AttendanceRecord.status == 'absent'
        ).all()

        absent_teachers = []
        for record, teacher, assignment, section in absent_records:
            absent_teachers.append({
                    'name': teacher.full_name,
                    'email': teacher.email,
                    'section': section.name,
                    'task_place': assignment.place_task,
                    'supervisor': plan.supervisor,
                    'warnings': teacher.warnings,
                    'notes': record.notes
                })

        # Send email to ibrahimfakhreyams@gmail.com with absent teachers report
        if absent_teachers:
            try:
                # Create email content
                email_content = f"""
                <h2>Daily Absent Teachers Report - {today.strftime('%Y-%m-%d')}</h2>
                <p>Total absent teachers: {len(absent_teachers)}</p>
                <table border="1" style="border-collapse: collapse; width: 100%;">
                    <tr style="background-color: #f2f2f2;">
                        <th style="padding: 8px;">Teacher Name</th>
                        <th style="padding: 8px;">Section</th>
                        <th style="padding: 8px;">Task/Place</th>
                        <th style="padding: 8px;">Supervisor</th>
                    </tr>
                """

                for teacher in absent_teachers:
                    email_content += f"""
                    <tr>
                        <td style="padding: 8px;">{teacher['name']}</td>
                        <td style="padding: 8px;">{teacher['section']}</td>
                        <td style="padding: 8px;">{teacher['task_place'] or 'N/A'}</td>
                        <td style="padding: 8px;">{teacher['supervisor'] or 'N/A'}</td>
                    </tr>
                    """

                email_content += "</table>"

                # Send email using existing email service
                import smtplib
                from email.mime.text import MIMEText
                from email.mime.multipart import MIMEMultipart

                # Multiple recipients
                recipients = ["islam.qamar@ams-benha.com", "ibrahimfakhreyams@gmail.com","yasser.alaraby@ams-benha.com","hadeer.tawfik@ams-benha.com","mohamad.mosalam@ams-benha.com"]

                msg = MIMEMultipart('alternative')
                msg['From'] = "amsprog2022@gmail.com"
                msg['To'] = ", ".join(recipients)
                msg['Subject'] = f"Daily Absent Teachers Report - {today.strftime('%Y-%m-%d')}"

                html_part = MIMEText(email_content, 'html')
                msg.attach(html_part)

                with smtplib.SMTP("smtp.gmail.com", 587) as server:
                    server.starttls()
                    server.login("amsprog2022@gmail.com", "xfleslznraphvqgc")
                    server.send_message(msg)

                flash(f'Email sent successfully to {", ".join(recipients)} with {len(absent_teachers)} absent teachers', 'success')

            except Exception as email_error:
                flash(f'Failed to send email: {str(email_error)}', 'error')
        else:
            flash('No absent teachers today - no email sent', 'info')

        # Render HTML page with absent teachers
        return render_template('absent_teachers_report.html',
                             absent_teachers=absent_teachers,
                             plan_date=today,
                             plan_exists=True,
                             total_absent=len(absent_teachers))

    except Exception as e:
        flash(f'Error generating report: {str(e)}', 'error')
        return render_template('absent_teachers_report.html',
                             absent_teachers=[],
                             plan_date=date.today(),
                             plan_exists=False)


@main_bp.route('/reports/absent-today')
@main_bp.route('/reports/absent/<date_str>')
def get_absent_people(date_str=None):
    """Get all absent people for a specific date (defaults to today)"""
    try:
        # Parse the date parameter or use today's date
        if date_str:
            try:
                target_date = datetime.strptime(date_str, '%Y-%m-%d').date()
            except ValueError:
                return jsonify({'error': 'Invalid date format. Use YYYY-MM-DD'}), 400
        else:
            target_date = date.today()

        # Query for all absent attendance records for the target date
        # Join with assignments to get teacher information
        query = db.session.query(AttendanceRecord).join(
            DutyAssignment, AttendanceRecord.assignment_id == DutyAssignment.id
        ).join(
            DutySection, DutyAssignment.section_id == DutySection.id
        ).join(
            DutyPlan, DutySection.plan_id == DutyPlan.id
        ).outerjoin(
            Teacher, DutyAssignment.teacher_id == Teacher.id
        ).filter(
            AttendanceRecord.date == target_date,
            AttendanceRecord.status == 'absent'
        ).order_by(
            DutyPlan.name,
            DutySection.order,
            DutyAssignment.order
        )

        absent_records = query.all()

        # Format the results
        absent_people = []
        for record in absent_records:
            assignment = record.assignment
            teacher_info = {
                'id': record.id,
                'date': record.date.isoformat(),
                'teacher_name': assignment.display_name,
                'teacher_email': assignment.teacher.email if assignment.teacher else None,
                'teacher_mobile': assignment.teacher.mobile if assignment.teacher else None,
                'section_name': assignment.section.name,
                'plan_name': assignment.section.plan.name,
                'day_of_week': assignment.section.plan.day_of_week,
                'place_task': assignment.place_task,
                'notes': record.notes,
                'is_placeholder': assignment.teacher is None
            }
            absent_people.append(teacher_info)

        return jsonify({
            'date': target_date.isoformat(),
            'total_absent': len(absent_people),
            'absent_people': absent_people
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@main_bp.route('/send-absent-notifications')
def send_absent_notifications():
    """
    Send individual email notifications to all absent teachers for today.
    This route manually triggers email notifications for today's absent teachers.
    """
    try:
        today = date.today()
        day_name = today.strftime("%A")

        # Get today's duty plan
        plan = DutyPlan.query.filter_by(day_of_week=day_name).first()

        if not plan:
            flash(f'No duty plan found for {day_name}', 'warning')
            return redirect(url_for('main.index'))

        # Query for absent teachers today
        absent_records = db.session.query(
            AttendanceRecord, Teacher, DutyAssignment, DutySection
        ).join(
            DutyAssignment, AttendanceRecord.assignment_id == DutyAssignment.id
        ).join(
            Teacher, DutyAssignment.teacher_id == Teacher.id
        ).join(
            DutySection, DutyAssignment.section_id == DutySection.id
        ).filter(
            AttendanceRecord.date == today,
            AttendanceRecord.status == 'absent',
            DutySection.plan_id == plan.id
        ).all()

        if not absent_records:
            flash('No absent teachers found for today', 'info')
            return redirect(url_for('main.index'))

        # Send individual notifications to each absent teacher
        sent_count = 0
        skipped_count = 0
        failed_count = 0

        for record, teacher, assignment, section in absent_records:
            try:
                result = send_attendance_notification(
                    teacher_email=teacher.email,
                    teacher_name=teacher.full_name,
                    status='absent',
                    day=day_name,
                    attendance_date=today,
                    section=section.name,
                    task_place=assignment.place_task,
                    supervisor=plan.supervisor,
                    warnings=teacher.warnings,
                    notes=record.notes
                )

                if result['success']:
                    if result.get('skipped', False):
                        skipped_count += 1
                    else:
                        sent_count += 1
                else:
                    failed_count += 1
                    flash(f'Failed to send email to {teacher.full_name}: {result["message"]}', 'error')

            except Exception as e:
                failed_count += 1
                flash(f'Error sending email to {teacher.full_name}: {str(e)}', 'error')

        # Show summary message
        total_teachers = len(absent_records)
        summary_parts = []

        if sent_count > 0:
            summary_parts.append(f'{sent_count} emails sent')
        if skipped_count > 0:
            summary_parts.append(f'{skipped_count} already sent (skipped)')
        if failed_count > 0:
            summary_parts.append(f'{failed_count} failed')

        summary = f'Notification summary for {total_teachers} absent teachers: {", ".join(summary_parts)}'

        if failed_count == 0:
            flash(summary, 'success')
        else:
            flash(summary, 'warning')

        return redirect(url_for('main.index'))

    except Exception as e:
        flash(f'Error processing absent teacher notifications: {str(e)}', 'error')
        return redirect(url_for('main.index'))


@main_bp.route('/api/delete-assignment', methods=['DELETE'])
def delete_assignment():
    """Delete a duty assignment from a specific day."""
    try:
        assignment_id = request.json.get('assignment_id')

        if not assignment_id:
            return jsonify({
                'success': False,
                'message': 'Assignment ID is required'
            }), 400

        # Find the assignment
        assignment = DutyAssignment.query.get(assignment_id)
        if not assignment:
            return jsonify({
                'success': False,
                'message': 'Assignment not found'
            }), 404

        # Get assignment details for response
        assignment_info = {
            'teacher_name': assignment.display_name,
            'section': assignment.section.name,
            'day': assignment.section.plan.day_of_week,
            'place_task': assignment.place_task
        }

        # Delete the assignment (cascade will handle attendance records)
        db.session.delete(assignment)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Successfully deleted assignment for {assignment_info["teacher_name"]}',
            'deleted_assignment': assignment_info
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': f'Error deleting assignment: {str(e)}'
        }), 500


@main_bp.route('/manage-assignments')
def manage_assignments():
    """Serve the manage assignments page with three forms."""
    return render_template('manage_assignments.html')


@main_bp.route('/api/search-assignments', methods=['GET'])
def search_assignments():
    """Search for duty assignments by day and query."""
    try:
        day = request.args.get('day', '').strip()
        query = request.args.get('q', '').strip()

        # Start with base query
        assignments_query = db.session.query(DutyAssignment).join(Teacher).join(DutySection).join(DutyPlan)

        # Filter by day if provided
        if day:
            assignments_query = assignments_query.filter(DutyPlan.day_of_week == day)

        # Filter by query if provided
        if query:
            assignments_query = assignments_query.filter(
                db.or_(
                    Teacher.full_name.ilike(f'%{query}%'),
                    DutySection.name.ilike(f'%{query}%'),
                    DutyAssignment.place_task.ilike(f'%{query}%')
                )
            )

        assignments = assignments_query.all()

        # Format results
        results = []
        for assignment in assignments:
            results.append({
                'id': assignment.id,
                'teacher_name': assignment.teacher.full_name,
                'day': assignment.section.plan.day_of_week,
                'section': assignment.section.name,
                'place_task': assignment.place_task,
                'order': assignment.order
            })

        return jsonify(results)

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error searching assignments: {str(e)}'
        }), 500


@main_bp.route('/api/edit-assignment', methods=['PUT'])
def edit_assignment():
    """Edit a duty assignment's role and positioning."""
    try:
        data = request.json
        assignment_id = data.get('assignment_id')
        new_section_name = data.get('section_name')
        new_place_task = data.get('place_task')
        new_order = data.get('order')

        if not assignment_id:
            return jsonify({
                'success': False,
                'message': 'Assignment ID is required'
            }), 400

        # Find the assignment
        assignment = DutyAssignment.query.get(assignment_id)
        if not assignment:
            return jsonify({
                'success': False,
                'message': 'Assignment not found'
            }), 404

        # Update section if provided
        if new_section_name:
            # Find or create the section in the same plan
            section = DutySection.query.filter_by(
                plan_id=assignment.section.plan_id,
                name=new_section_name
            ).first()

            if not section:
                # Create new section
                max_order = db.session.query(db.func.max(DutySection.order)).filter_by(
                    plan_id=assignment.section.plan_id
                ).scalar() or 0

                section = DutySection(
                    plan_id=assignment.section.plan_id,
                    name=new_section_name,
                    order=max_order + 1
                )
                db.session.add(section)
                db.session.flush()

            assignment.section_id = section.id

        # Update place_task if provided
        if new_place_task is not None:
            assignment.place_task = new_place_task

        # Update order if provided
        if new_order is not None:
            assignment.order = new_order

        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Successfully updated assignment for {assignment.display_name}',
            'updated_assignment': {
                'id': assignment.id,
                'teacher_name': assignment.display_name,
                'section': assignment.section.name,
                'place_task': assignment.place_task,
                'order': assignment.order,
                'day': assignment.section.plan.day_of_week
            }
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': f'Error updating assignment: {str(e)}'
        }), 500


@main_bp.route('/api/add-assignment', methods=['POST'])
def add_assignment():
    """Add new members to a specific day with their roles."""
    try:
        data = request.json
        day_of_week = data.get('day_of_week')
        teacher_name = data.get('teacher_name')
        teacher_email = data.get('teacher_email')
        section_name = data.get('section_name')
        place_task = data.get('place_task')

        if not day_of_week or not teacher_name or not section_name:
            return jsonify({
                'success': False,
                'message': 'Day of week, teacher name, and section name are required'
            }), 400

        # Validate day format
        day_of_week = day_of_week.capitalize()
        valid_days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
        if day_of_week not in valid_days:
            return jsonify({
                'success': False,
                'message': f'Invalid day. Must be one of: {", ".join(valid_days)}'
            }), 400

        # Find the duty plan for the specified day
        plan = DutyPlan.query.filter_by(day_of_week=day_of_week).first()
        if not plan:
            return jsonify({
                'success': False,
                'message': f'No duty plan found for {day_of_week}'
            }), 404

        # Try to find existing teacher by name or email
        teacher = None
        if teacher_email:
            teacher = Teacher.query.filter_by(email=teacher_email).first()

        if not teacher and teacher_name:
            teacher = Teacher.query.filter_by(full_name=teacher_name).first()

        # Find or create the section
        section = DutySection.query.filter_by(
            plan_id=plan.id,
            name=section_name
        ).first()

        if not section:
            # Create new section
            max_order = db.session.query(db.func.max(DutySection.order)).filter_by(
                plan_id=plan.id
            ).scalar() or 0

            section = DutySection(
                plan_id=plan.id,
                name=section_name,
                order=max_order + 1
            )
            db.session.add(section)
            db.session.flush()

        # Get next order for assignment in this section
        max_assignment_order = db.session.query(db.func.max(DutyAssignment.order)).filter_by(
            section_id=section.id
        ).scalar() or 0

        # Create new assignment
        assignment = DutyAssignment(
            section_id=section.id,
            teacher_id=teacher.id if teacher else None,
            placeholder_name=teacher_name if not teacher else None,
            order=max_assignment_order + 1,
            place_task=place_task
        )

        db.session.add(assignment)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Successfully added {teacher_name} to {day_of_week} plan',
            'new_assignment': {
                'id': assignment.id,
                'teacher_name': assignment.display_name,
                'teacher_email': teacher.email if teacher else teacher_email,
                'section': section.name,
                'place_task': assignment.place_task,
                'order': assignment.order,
                'day': day_of_week,
                'is_existing_teacher': teacher is not None
            }
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': f'Error adding assignment: {str(e)}'
        }), 500
